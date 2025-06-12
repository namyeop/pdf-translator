import streamlit as st
import random
import time
import os

# Langchain imports
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate, PromptTemplate  # MODIFIED
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough  # For passing input through
from langchain_community.document_loaders import (
    PyPDFLoader,
)  # For loading PDF documents
from langchain_text_splitters import (
    RecursiveCharacterTextSplitter,
)  # For splitting text into chunks
from pdf2docx import Converter  # Added for PDF to DOCX conversion
from docx import Document  # ADDED


def load_and_chunk_documents(docs_path: str) -> list:
    """
    Loads documents from the specified directory, processes PDF files,
    and splits them into manageable chunks.
    """
    all_chunks = []
    if not os.path.isdir(docs_path):
        st.error(f"Docs directory not found: {docs_path}")
        return all_chunks

    try:
        doc_files = [f for f in os.listdir(docs_path) if f.lower().endswith(".pdf")]
        if not doc_files:
            st.info(f"No PDF documents found in {docs_path}")
            return all_chunks

        loaded_documents = []
        for doc_file in doc_files:
            file_path = os.path.join(docs_path, doc_file)
            try:
                loader = PyPDFLoader(file_path)
                loaded_documents.extend(
                    loader.load()
                )  # PyPDFLoader loads pages as individual documents
                st.write(f"Loaded {doc_file}")
            except Exception as e:
                st.error(f"Error loading {doc_file}: {e}")

        if not loaded_documents:
            st.warning("No documents were successfully loaded.")
            return all_chunks

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,  # Size of each chunk
            chunk_overlap=200,  # Overlap between chunks
            length_function=len,
        )
        all_chunks = text_splitter.split_documents(loaded_documents)
        st.success(
            f"Successfully loaded and chunked {len(loaded_documents)} page(s) into {len(all_chunks)} chunks."
        )

    except Exception as e:
        st.error(f"An error occurred during document loading and chunking: {e}")

    return all_chunks


def convert_pdf_to_docx(pdf_file_path: str, docx_file_path: str) -> bool:
    """
    Converts a PDF file to a DOCX file using pdf2docx.

    Args:
        pdf_file_path: Path to the input PDF file.
        docx_file_path: Path to save the output DOCX file.

    Returns:
        True if conversion was successful, False otherwise.
    """
    if not os.path.exists(pdf_file_path):
        st.error(f"PDF file not found: {pdf_file_path}")
        return False

    try:
        st.write(f"Converting {os.path.basename(pdf_file_path)} to DOCX...")
        # Initialize Converter
        cv = Converter(pdf_file_path)
        # Convert to DOCX
        cv.convert(docx_file_path, start=0, end=None)  # Converts all pages
        # Close the converter
        cv.close()
        st.success(f"Successfully converted PDF to DOCX: {docx_file_path}")
        return True
    except Exception as e:
        st.error(f"Error during PDF to DOCX conversion: {e}")
        return False


def read_docx_content(docx_file_path: str) -> str:
    """Reads text content from a DOCX file."""
    if not os.path.exists(docx_file_path):
        st.error(f"DOCX file not found: {docx_file_path}")
        return ""
    try:
        doc = Document(docx_file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return "\\n".join(full_text)
    except Exception as e:
        st.error(f"Error reading DOCX file {docx_file_path}: {e}")
        return ""


def advanced_replace_docx_translation(
    original_docx_path: str, translated_text: str, output_path: str
) -> bool:
    """Advanced method to replace text while preserving complete document structure."""
    try:
        import copy
        from docx.shared import RGBColor

        # Load the original document
        doc = Document(original_docx_path)

        # Extract all text content with structure preservation
        all_paragraphs_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                all_paragraphs_text.append(paragraph.text.strip())

        # Create a more intelligent mapping
        original_full_text = " ".join(all_paragraphs_text)

        # Split translated text more intelligently
        translated_lines = [
            line.strip() for line in translated_text.split("\n") if line.strip()
        ]

        # Try to match paragraph count
        if len(translated_lines) < len(all_paragraphs_text):
            # If we have fewer translated paragraphs, try to split longer ones
            while len(translated_lines) < len(all_paragraphs_text) and translated_lines:
                longest_line_idx = max(
                    range(len(translated_lines)), key=lambda i: len(translated_lines[i])
                )
                longest_line = translated_lines[longest_line_idx]
                if ". " in longest_line:
                    parts = longest_line.split(". ", 1)
                    translated_lines[longest_line_idx] = parts[0] + "."
                    translated_lines.insert(longest_line_idx + 1, parts[1])
                else:
                    break

        # Replace text paragraph by paragraph while preserving formatting
        para_index = 0
        for paragraph in doc.paragraphs:
            if paragraph.text.strip() and para_index < len(translated_lines):
                # Store all formatting information before modification
                runs_info = []
                for run in paragraph.runs:
                    run_info = {
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                        "font_name": run.font.name,
                        "font_size": run.font.size,
                        "color": run.font.color.rgb if run.font.color.rgb else None,
                    }
                    runs_info.append(run_info)

                # Clear paragraph but keep style
                paragraph_style = paragraph.style
                paragraph_alignment = paragraph.alignment
                paragraph.clear()

                # Add translated text with original formatting
                if runs_info:
                    # Use the formatting from the first run for the entire translated text
                    first_run_info = runs_info[0]
                    new_run = paragraph.add_run(translated_lines[para_index])

                    # Apply formatting
                    if first_run_info["bold"] is not None:
                        new_run.bold = first_run_info["bold"]
                    if first_run_info["italic"] is not None:
                        new_run.italic = first_run_info["italic"]
                    if first_run_info["underline"] is not None:
                        new_run.underline = first_run_info["underline"]
                    if first_run_info["font_name"]:
                        new_run.font.name = first_run_info["font_name"]
                    if first_run_info["font_size"]:
                        new_run.font.size = first_run_info["font_size"]
                    if first_run_info["color"]:
                        new_run.font.color.rgb = first_run_info["color"]
                else:
                    # Fallback: just add the text
                    paragraph.add_run(translated_lines[para_index])

                # Restore paragraph-level formatting
                paragraph.style = paragraph_style
                paragraph.alignment = paragraph_alignment

                para_index += 1

        doc.save(output_path)
        return True
    except Exception as e:
        st.error(f"Error in advanced DOCX translation replacement: {e}")
        return False


def translate_text_with_llm(
    text_to_translate: str,
    target_language: str = "Italian",  # MODIFIED default for consistency
) -> str:
    """Translates text using the configured LLM."""
    global llm  # Ensure llm is accessible
    st.info(
        f"translate_text_with_llm 호출됨. 대상 언어: {target_language}"
    )  # ADDED FOR DEBUG

    if not llm:
        st.error("LLM not available for translation.")
        return "Translation service not available (LLM not initialized)."  # MODIFIED for clarity
    if not text_to_translate.strip():
        st.warning("번역할 내용이 없습니다.")  # ADDED FOR DEBUG
        return "No text provided for translation."

    translation_prompt_template = PromptTemplate.from_template(
        "Translate the following text into {target_language}. "
        "IMPORTANT: Preserve the exact paragraph structure and line breaks. "
        "Each paragraph should remain as a separate paragraph in the translation. "
        "Provide only the translated text without any introductory phrases or explanations.\n\n"
        "Text: {text}\n\n"
        "Translation:"
    )

    translation_chain = translation_prompt_template | llm | StrOutputParser()

    try:
        st.info(
            f"LLM으로 번역 요청 중... 내용 일부 (최대 200자): {text_to_translate[:200]}..."
        )  # ADDED FOR DEBUG
        # st.write(f"Translating text to {target_language}...") # Original st.write, can be kept or removed
        translated_text = translation_chain.invoke(
            {"text": text_to_translate, "target_language": target_language}
        )
        st.info(
            f"LLM으로부터 받은 번역 일부 (최대 200자): {translated_text[:200]}..."
        )  # ADDED FOR DEBUG

        if not translated_text or not translated_text.strip():
            st.warning(
                f"LLM이 {target_language}로 빈 번역을 반환했습니다."
            )  # ADDED warning
            return "번역 결과가 비어있습니다."  # MODIFIED to return a message
        return translated_text
    except Exception as e:
        st.error(f"번역 중 오류 발생: {e}")  # Original error handling
        return f"Translation failed: {e}"


# Determine the absolute path to the docs directory
# Assumes this script (app.py) is in the app's root directory and 'docs' is a subdirectory. # MODIFIED
APP_DIR = os.path.dirname(os.path.abspath(__file__))
DOCS_DIR = os.path.join(APP_DIR, "docs")

# Load documents when the app starts
chunked_documents = load_and_chunk_documents(DOCS_DIR)

st.title("✨ Chat with RAG ✨")  # Updated title

# --- PDF to DOCX Converter UI ---

# Langsmith & OpenAI API Key Setup
# Using Streamlit secrets (recommended for deployment)
# Ensure these secrets are set in your Streamlit Cloud config or a .streamlit/secrets.toml file
try:
    os.environ["LANGCHAIN_TRACING_V2"] = "true"  # Enable Langsmith tracing

    langchain_api_key = st.secrets.get("LANGCHAIN_API_KEY")
    if langchain_api_key:
        os.environ["LANGCHAIN_API_KEY"] = langchain_api_key
    else:
        st.warning(
            "LANGCHAIN_API_KEY not found in Streamlit secrets. Langsmith tracing might not work."
        )

    langchain_project = st.secrets.get(
        "LANGCHAIN_PROJECT", "Streamlit RAG Chat"
    )  # Default project name
    os.environ["LANGCHAIN_PROJECT"] = langchain_project

    openai_api_key = st.secrets.get("OPENAI_API_KEY")
    if openai_api_key:
        os.environ["OPENAI_API_KEY"] = openai_api_key
        llm = ChatOpenAI(model="gpt-4.1-2025-04-14")
    else:
        st.error(
            "OPENAI_API_KEY not found in Streamlit secrets. RAG and Translation functionality requires an OpenAI API key."
        )
        llm = None

except (
    AttributeError,
    NameError,
):  # Handles cases where st.secrets might not be available
    st.warning(
        "Streamlit secrets (st.secrets) not available. Attempting to use environment variables directly."
    )
    llm = None  # Initialize llm to None before checking environment variables
    if not os.getenv("LANGCHAIN_API_KEY"):
        st.warning("LANGCHAIN_API_KEY environment variable not set.")
    if not os.getenv("OPENAI_API_KEY"):
        st.error(
            "OPENAI_API_KEY environment variable not set. RAG and Translation functionality requires an OpenAI API key."
        )
    else:
        llm = ChatOpenAI(model="gpt-4.1-2025-04-14")
    if not os.getenv("LANGCHAIN_PROJECT"):
        os.environ["LANGCHAIN_PROJECT"] = "Streamlit RAG Chat (direct env)"

# --- PDF Translation and DOCX Export ---
if "pdf_process_data" not in st.session_state:
    st.session_state.pdf_process_data = {
        "uploaded_filename": None,
        "docx_path": None,
        "docx_filename": None,
        "status_message": None,
        "translated_text_for_chat": None,
        "processed": False,
    }
if (
    "messages" not in st.session_state
):  # Ensure messages is initialized for chat app parts
    st.session_state.messages = []


new_pdf_file = st.file_uploader(
    "Upload PDF for Translation & DOCX Export", type="pdf", key="pdf_upload_main"
)

if new_pdf_file is not None:
    if (
        st.session_state.pdf_process_data["uploaded_filename"] != new_pdf_file.name
        or not st.session_state.pdf_process_data["processed"]
    ):

        st.session_state.pdf_process_data = {
            "uploaded_filename": new_pdf_file.name,
            "docx_path": None,
            "docx_filename": None,
            "status_message": "Initializing processing...",
            "translated_text_for_chat": None,
            "processed": False,
        }
        temp_pdf_path = os.path.join(APP_DIR, new_pdf_file.name)
        with open(temp_pdf_path, "wb") as f:
            f.write(new_pdf_file.getbuffer())

        docx_filename_out = os.path.splitext(new_pdf_file.name)[0] + ".docx"
        temp_docx_path_out = os.path.join(APP_DIR, docx_filename_out)

        st.session_state.pdf_process_data["docx_filename"] = docx_filename_out
        st.session_state.pdf_process_data["docx_path"] = temp_docx_path_out

        current_status_messages = []

        with st.spinner(
            f"Processing {new_pdf_file.name}... (Convert, Read, Translate)"
        ):
            conversion_ok = convert_pdf_to_docx(temp_pdf_path, temp_docx_path_out)

            if conversion_ok:
                current_status_messages.append(
                    f"Successfully converted to {docx_filename_out}."
                )
                content = read_docx_content(temp_docx_path_out)
                if content:
                    current_status_messages.append("Content read from DOCX.")
                    translated = translate_text_with_llm(content, "Italian")
                    st.session_state.pdf_process_data["translated_text_for_chat"] = (
                        translated
                    )
                    current_status_messages.append("Translation complete.")

                    # Create new DOCX with translated content while preserving original formatting
                    translated_docx_filename = f"translated_{docx_filename_out}"
                    translated_docx_path = os.path.join(
                        APP_DIR, translated_docx_filename
                    )

                    # Replace content in original DOCX with translated text while preserving formatting
                    docx_created = advanced_replace_docx_translation(
                        temp_docx_path_out, translated, translated_docx_path
                    )
                    if docx_created:
                        # Update session state with new translated DOCX path
                        st.session_state.pdf_process_data["docx_path"] = (
                            translated_docx_path
                        )
                        st.session_state.pdf_process_data["docx_filename"] = (
                            translated_docx_filename
                        )
                        current_status_messages.append(
                            "Translated DOCX created with original formatting."
                        )

                        # Clean up the original temporary DOCX
                        if os.path.exists(temp_docx_path_out):
                            try:
                                os.remove(temp_docx_path_out)
                            except Exception as e:
                                st.warning(f"Could not remove original DOCX: {e}")
                    else:
                        current_status_messages.append(
                            "Failed to create translated DOCX."
                        )

                    st.session_state.messages.append(
                        {
                            "role": "assistant",
                            "content": f"번역된 내용 ({new_pdf_file.name}):\n\n{translated}",
                        }
                    )
                else:
                    err_msg = "Could not read content from DOCX."
                    current_status_messages.append(err_msg)
                    st.session_state.pdf_process_data["translated_text_for_chat"] = (
                        err_msg
                    )
            else:
                err_msg = "PDF to DOCX conversion failed."
                current_status_messages.append(err_msg)
                st.session_state.pdf_process_data["translated_text_for_chat"] = err_msg

        st.session_state.pdf_process_data["status_message"] = " ".join(
            current_status_messages
        )
        st.session_state.pdf_process_data["processed"] = True

        if os.path.exists(temp_pdf_path):
            try:
                os.remove(temp_pdf_path)
                # st.info(f"Removed temporary PDF: {temp_pdf_path}") # Optional: can be noisy
            except Exception as e:
                st.warning(f"Could not remove temporary PDF {temp_pdf_path}: {e}")

        st.rerun()

if (
    st.session_state.pdf_process_data["processed"]
    and st.session_state.pdf_process_data["uploaded_filename"]
):
    status_msg = st.session_state.pdf_process_data["status_message"]
    if "failed" in status_msg.lower() or "could not read" in status_msg.lower():
        st.error(status_msg)
    else:
        st.success(status_msg)

    if st.session_state.pdf_process_data["docx_path"] and os.path.exists(
        st.session_state.pdf_process_data["docx_path"]
    ):
        try:
            with open(st.session_state.pdf_process_data["docx_path"], "rb") as fp:
                st.download_button(
                    label=f"Download Translated {st.session_state.pdf_process_data['docx_filename']}",
                    data=fp,
                    file_name=st.session_state.pdf_process_data["docx_filename"],
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="docx_download_main",
                )
        except Exception as e:
            st.error(f"Error making DOCX download available: {e}")

if st.session_state.pdf_process_data["uploaded_filename"] is not None:
    if st.button(
        "Clear processed PDF and upload another", key="clear_processed_pdf_btn"
    ):
        if st.session_state.pdf_process_data["docx_path"] and os.path.exists(
            st.session_state.pdf_process_data["docx_path"]
        ):
            try:
                os.remove(st.session_state.pdf_process_data["docx_path"])
                st.info(
                    f"Removed DOCX file: {st.session_state.pdf_process_data['docx_filename']}"
                )
            except Exception as e:
                st.warning(f"Could not remove DOCX file: {e}")

        st.session_state.pdf_process_data = {
            "uploaded_filename": None,
            "docx_path": None,
            "docx_filename": None,
            "status_message": None,
            "translated_text_for_chat": None,
            "processed": False,
        }
        st.rerun()
# --- End PDF Translation and DOCX Export ---


# --- RAG Setup ---
# The old get_dummy_retriever is no longer needed.

# Prompt Template for RAG
rag_prompt_template = ChatPromptTemplate.from_template(
    """Answer the following question based only on the provided context. If the context doesn't contain the answer, say you don't know.

Context:
{context}

Question: {question}
"""
)  # Simplified RAG Chain


def retrieve_context_from_chunks(query: str, chunks: list):
    if not chunks:
        return "No documents have been loaded or processed. Cannot retrieve context."
    # For now, concatenate all document chunks as context.
    # This is a placeholder for a more sophisticated retrieval (e.g., vector search).
    # Consider limiting the amount of context if it becomes too large.
    context = "\n\n---\n\n".join([doc.page_content for doc in chunks])
    return context


if llm:
    if chunked_documents:
        rag_chain = (
            RunnablePassthrough.assign(
                context=lambda x: retrieve_context_from_chunks(
                    x["question"], chunked_documents
                )
            )
            | rag_prompt_template
            | llm
            | StrOutputParser()
        )
    else:
        st.error(
            "Documents could not be loaded. RAG chain will not use document context."
        )
        # Fallback: Initialize rag_chain to None or a version that doesn't rely on docs
        # For now, if docs fail, RAG_response_generator will indicate issues.
        rag_chain = None
else:
    rag_chain = None  # LLM is not initialized

# --- End RAG Setup ---


# Initialize chat history
if "messages" not in st.session_state:
    st.session_state.messages = []

# Display chat messages from history on app rerun
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])


def RAG_response_generator(user_input: str):
    if not llm or not rag_chain:
        yield "RAG chain is not configured (likely due to missing API keys or LLM initialization failure). "
        return

    try:
        # The input to the chain should be a dictionary
        input_dict = {"question": user_input}
        full_response_parts = []
        for chunk in rag_chain.stream(input_dict):
            yield chunk
            full_response_parts.append(chunk)
            # time.sleep(0.02) # Optional: small delay for typing effect, adjust as needed
        # The full response is now "".join(full_response_parts)
        # This will be handled by st.write_stream which returns the full string.
    except Exception as e:
        st.error(f"Error during RAG generation: {e}")
        yield f"Sorry, I encountered an error while generating a response: {str(e)}. "


# React to user input
if prompt := st.chat_input("메시지를 입력하세요..."):
    # Add user message to chat history
    st.session_state.messages.append({"role": "user", "content": prompt})
    # Display user message in chat message container
    with st.chat_message("user"):
        st.markdown(prompt)

    # Display assistant response in chat message container
    with st.chat_message("assistant"):
        if llm and rag_chain:  # Check if LLM and chain are ready
            full_response = st.write_stream(RAG_response_generator(prompt))
            st.session_state.messages.append(
                {"role": "assistant", "content": full_response}
            )
        else:
            # Fallback if LLM/RAG is not set up
            fallback_message = "RAG system not fully configured. Please check API key settings and ensure the LLM is available."
            st.markdown(fallback_message)
            st.session_state.messages.append(
                {"role": "assistant", "content": fallback_message}
            )
